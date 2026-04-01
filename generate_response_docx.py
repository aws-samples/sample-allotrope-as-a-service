"""Convert AWS Response to Merck Review to Word document."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        shading = cell._element.get_or_add_tcPr()
        shading.append(shading.makeelement(qn('w:shd'), {qn('w:fill'): '232F3E', qn('w:val'): 'clear'}))
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def build():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    t = doc.add_heading('AWS Response to Merck ASM Conversion Review', level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('In response to: 2026-03-17 AWS ASM Conversion Review & Merck ASM Basic Design Principles v33')
    doc.add_paragraph()

    # ── Q1 ──
    doc.add_heading('Responses to Questions', level=1)
    doc.add_heading('Question 1: How can the platform capture instrument-specific metadata?', level=2)
    doc.add_paragraph(
        'The instrument config file (JSON) is the right vehicle. We will extend the schema to include '
        'metadata fields that capture the information currently derived from the mount path: geography, '
        'site, UNC path, and timezone.'
    )
    doc.add_paragraph(
        'The timezone field solves the issue where the July file (European) was interpreted with US date format. '
        'The UNC path field provides traceability back to the source file server. '
        'The instrument config is created once per instrument and reused for every file.'
    )
    doc.add_paragraph(
        'For the future Greengrass edge client deployment, the mount path metadata would be captured '
        'automatically since the edge client knows which instrument folder it is watching.'
    )

    # ── Q2 ──
    doc.add_heading('Question 2: Can the AI agent provide a confidence level for each mapping?', level=2)
    doc.add_paragraph('Yes, but with an important distinction between conversion routes:')
    p = doc.add_paragraph()
    p.add_run('Custom converters (Route A): ').bold = True
    p.add_run('Confidence is 100% by design. Every field mapping is deterministic. The Data Integrity Verification report proves this field-by-field.')
    p = doc.add_paragraph()
    p.add_run('Allotropy library (Route B): ').bold = True
    p.add_run('Confidence is 100% for supported instruments. Rule-based parsers, no AI inference.')
    p = doc.add_paragraph()
    p.add_run('ATaaS / AI-powered (Route C): ').bold = True
    p.add_run('This is where per-mapping confidence levels are valuable. We can enhance ATaaS to provide High/Medium/Low confidence per field mapping.')
    doc.add_paragraph(
        'We recommend custom converters or allotropy for production use where confidence is deterministic. '
        'ATaaS is best suited for initial analysis of unknown formats with human-in-the-loop review.'
    )

    # ── Q3 ──
    doc.add_heading('Question 3: Which validation levels are we using today?', level=2)
    styled_table(doc,
        ['Merck Level', 'Description', 'Our Equivalent', 'Status'],
        [
            ['Level 1', 'ASM schema validation (JSON + ASM tags)', 'DVaaS with allotropy validation script', 'Implemented'],
            ['Level 2', 'Data quality assessment (ASM vs raw data)', 'Data Integrity Verification (field_mapping)', 'Work in progress \u2014 Nova FLEX2 done, rolling out'],
            ['Level 3', 'Allotrope ASM building principles', 'Not yet implemented', 'Planned \u2014 add as DVaaS rules'],
            ['Level 4', 'Merck-specific principles', 'Not yet implemented', 'Planned \u2014 configurable per-customer rules'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Our Data Integrity Verification directly addresses your Level 2 pain point. You noted that Level 2 requires '
        '"significant instrument-specific coding." Our approach solves this: the converter emits the field_mapping '
        'during conversion, so verification is automatic. No TOSCA database, no SQL queries, no instrument-specific scripts.'
    )
    doc.add_paragraph(
        'We will update DVaaS to tag each finding with the corresponding validation level (L1/L2/L3/L4) '
        'so the human-in-the-loop reviewer can prioritize accordingly.'
    )

    # ── Remediation ──
    doc.add_heading('Remediation Plan', level=1)
    doc.add_heading('Issue Tracker', level=2)
    styled_table(doc,
        ['#', 'Category', 'Issue', 'Severity', 'Status'],
        [
            ['1', '[META]', 'Measurement time missing timezone/geography context', 'Medium', 'Open'],
            ['2', '[ERROR]', 'Sample type uses invalid ASM term', 'High', 'Open'],
            ['3', '[ERROR]', 'Skipped processed data: total cell count, viable cell count, cell density dilution factor', 'High', 'Open'],
            ['4', '[ERROR]', 'Cell density dilution factor not nested under data processing document', 'Medium', 'Open'],
            ['5', '[ERROR][RULE]', 'Calculated fields were ignored', 'High', 'Fixed'],
            ['6', '[ERROR][RULE]', 'Custom fields were ignored', 'High', 'Fixed'],
            ['7', '[CONFIG]', 'Missing Data System Document', 'High', 'Partially fixed \u2014 needs UNC path'],
            ['8', '[CONFIG]', 'Missing Device Document with device type', 'High', 'Fixed'],
            ['9', '[META]', 'No traceability to original source file (UNC path)', 'Medium', 'Open'],
            ['10', '[RULE]', 'Custom fields have manually added units not in source', 'Medium', 'Open'],
            ['11', '[RULE]', 'Missing @index for array ordering', 'Low', 'Open'],
            ['12', '[RULE]', 'Calculated data missing per-calculation data source linkage', 'Medium', 'Open'],
            ['13', '[RULE]', 'No processed data aggregate document layer', 'Medium', 'Open'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('Already Fixed Since Review', level=2)
    fixes = [
        ('Issue 5 \u2014 Calculated fields:', 'Converter now includes calculated data aggregate document with temperature-corrected pH, PO2, PCO2, and bicarbonate.'),
        ('Issue 6 \u2014 Custom fields:', 'Converter now maps all custom fields to custom information aggregate document including lot numbers, flow times, vessel pressure, dilution ratio, and all string metadata.'),
        ('Issue 8 \u2014 Device Document:', 'Converter includes device system document with device identifier, product manufacturer, and device type = "solution analyzer".'),
    ]
    for title, desc in fixes:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(' ' + desc)

    doc.add_heading('Phase 1 \u2014 High Priority', level=2)
    styled_table(doc,
        ['#', 'Issue', 'Effort', 'Change'],
        [
            ['1', 'Timezone in instrument config', '1 hour', 'Add location.timezone to config, use in timestamp parsing'],
            ['2', 'Invalid sample type term', '30 min', 'Map to valid ASM term per schema'],
            ['9', 'UNC path traceability', '1 hour', 'Read from instrument config, write to data system document'],
            ['10', 'Remove manually added units from custom fields', '1 hour', 'Only include units when present in source data'],
            ['12', 'Per-calculation data source linkage', '1 hour', 'Nest data source aggregate inside each calculated data document'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('Phase 2 \u2014 Medium Priority', level=2)
    styled_table(doc,
        ['#', 'Issue', 'Effort', 'Change'],
        [
            ['3', 'Processed data (cell counts, dilution factor)', '2 hours', 'Add processed data aggregate document'],
            ['4', 'Dilution factor nesting', '30 min', 'Nest under data processing document'],
            ['11', '@index for array ordering', '1 hour', 'Add @index based on source order or timestamp'],
            ['13', 'Processed data layer', '2 hours', 'Distinguish raw measurement vs instrument transformation'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('Phase 3 \u2014 Validation Enhancement', level=2)
    styled_table(doc,
        ['Item', 'Effort', 'Change'],
        [
            ['DVaaS Level 3 rules (Allotrope principles)', '4 hours', 'Add design principle checks to validation'],
            ['DVaaS Level 4 rules (Merck-specific)', '4 hours', 'Add configurable customer-specific rules'],
            ['Validation level labels on findings', '2 hours', 'Tag each error/warning with L1/L2/L3/L4'],
            ['ATaaS confidence scores per mapping', '4 hours', 'Add per-field confidence to AI-generated mappings'],
        ]
    )
    doc.add_paragraph()

    # ── Level 2 comparison ──
    doc.add_heading('Data Integrity Verification \u2014 Addressing Level 2', level=1)
    doc.add_paragraph(
        'Your Level 2 validation (TOSCA-based source-to-ASM comparison) requires significant per-instrument coding. '
        'Our Data Integrity Verification provides the same assurance with a fundamentally different approach:'
    )
    styled_table(doc,
        ['Aspect', 'Merck Level 2 (TOSCA)', 'AWS Data Integrity Verification'],
        [
            ['Approach', 'Post-conversion comparison via SQL', 'Converter emits field_mapping during conversion'],
            ['Per-instrument effort', 'Significant \u2014 SQLite, SQL queries, tolerances', 'Zero \u2014 field_mapping built into converter'],
            ['Tolerance handling', 'SQL-based rounding and TOSCA parameters', 'Exact match \u2014 no rounding needed'],
            ['Coverage', 'Depends on SQL queries written', 'Every field the converter touches'],
            ['Output', 'TOSCA report', 'Dashboard table + API response'],
            ['Maintenance', 'Update SQL when converter changes', 'Automatic \u2014 updates with converter'],
        ]
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Key advantage: ').bold = True
    p.add_run(
        'Because the converter emits the mapping at conversion time, there is no separate comparison step, '
        'no database, and no instrument-specific SQL. The integrity proof is a byproduct of the conversion itself.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Current status: Implemented for Nova FLEX2 (all 40 CSV columns). Required for all custom converters '
        'per our requirements spec. Planned for allotropy-based conversions.'
    )

    # ── Alignment ──
    doc.add_heading('Alignment with Merck Design Principles', level=1)
    doc.add_heading('What We Align On Today', level=2)
    styled_table(doc,
        ['Principle', 'Our Implementation', 'Status'],
        [
            ['Instrument hardcoded terms', 'Instrument config JSON file', 'Aligned \u2014 needs UNC path and timezone'],
            ['Data hierarchical structure', 'Follow solution-analyzer schema', 'Aligned'],
            ['ASM terms lowercase with spaces', 'Terms match schema (pO2, pCO2, pH, osmolality)', 'Aligned'],
            ['Custom fields 1:1 transfer', 'Custom information aggregate document', 'Aligned \u2014 need to remove manual units'],
            ['Calculated fields with traceability', 'Calculated data aggregate document', 'Aligned \u2014 need to fix nesting'],
            ['Manufacturer in lowercase', '"nova biomedical"', 'Aligned'],
            ['Source file traceability', 'Data system document with file name, converter info', 'Partial \u2014 need UNC path'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('What We Need to Add', level=2)
    styled_table(doc,
        ['Principle', 'Gap', 'Plan'],
        [
            ['Processed data aggregate document', 'Not implemented', 'Phase 2'],
            ['@index for array ordering', 'Not implemented', 'Phase 2'],
            ['HarmonizingBaseTerms integration', 'Not formalized', 'Accept as input alongside instrument config'],
            ['Custom field units only from source', 'We add units manually', 'Phase 1'],
            ['Statistic datum role', 'Not implemented', 'Future'],
        ]
    )
    doc.add_paragraph()

    # ── Next Steps ──
    doc.add_heading('Next Steps', level=1)
    steps = [
        'Share this response with Merck for review',
        'Phase 1: Implement high-priority fixes (timezone, sample type, UNC path, custom field units, data source nesting) \u2014 estimated 5 hours',
        'Phase 2: Implement medium-priority fixes (processed data, @index, dilution factor nesting) \u2014 estimated 6 hours',
        'Phase 3: Validation enhancement (L3/L4 rules, level labels, confidence scores) \u2014 estimated 14 hours',
        'Ongoing: Accept and integrate HarmonizingBaseTerms spreadsheet as converter configuration input',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {s}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Built with: ').bold = True
    p.add_run('AWS Lambda, AWS Bedrock Claude, Allotropy Library, React, AWS CDK')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


if __name__ == '__main__':
    doc = build()
    path = 'docs/AWS-Response-to-Merck-Review.docx'
    doc.save(path)
    print(f'Saved to {path}')
